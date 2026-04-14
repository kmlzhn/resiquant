"""Extract raw text from PDF, XLSX, XLS, DOCX files."""
from __future__ import annotations
import io
import warnings
from dataclasses import dataclass
from pathlib import Path

warnings.filterwarnings("ignore")


@dataclass
class ParsedDocument:
    name: str
    pages: list[str]  # one entry per page / sheet

    @property
    def full_text(self) -> str:
        return "\n\n".join(self.pages)


def parse_file(path: Path | str, content: bytes | None = None) -> ParsedDocument:
    """Parse a file from disk path or raw bytes. Returns ParsedDocument."""
    path = Path(path)
    suffix = path.suffix.lower()

    if content is None:
        content = path.read_bytes()

    if suffix == ".pdf":
        return _parse_pdf(path.name, content)
    elif suffix in (".xlsx", ".xls"):
        return _parse_excel(path.name, content, suffix)
    elif suffix == ".docx":
        return _parse_docx(path.name, content)
    else:
        # Try as plain text
        try:
            return ParsedDocument(name=path.name, pages=[content.decode("utf-8", errors="replace")])
        except Exception:
            return ParsedDocument(name=path.name, pages=["[unreadable file]"])


def _parse_pdf(name: str, content: bytes) -> ParsedDocument:
    import pdfplumber
    pages: list[str] = []
    with pdfplumber.open(io.BytesIO(content)) as pdf:
        for i, page in enumerate(pdf.pages):
            text = page.extract_text() or ""
            if text.strip():
                pages.append(f"[Page {i+1}]\n{text.strip()}")
    if not pages:
        pages = ["[PDF contained no extractable text]"]
    return ParsedDocument(name=name, pages=pages)


def _parse_excel(name: str, content: bytes, suffix: str) -> ParsedDocument:
    import pandas as pd
    engine = "openpyxl" if suffix == ".xlsx" else "xlrd"
    sheets: list[str] = []
    try:
        xls = pd.read_excel(io.BytesIO(content), sheet_name=None, engine=engine, dtype=str, header=None)
        for sheet_name, df in xls.items():
            df = df.dropna(how="all")
            rows = []
            for _, row in df.iterrows():
                cells = [str(v).strip() for v in row if str(v).strip() not in ("", "nan")]
                if cells:
                    rows.append(" | ".join(cells))
            text = f"[Sheet: {sheet_name}]\n" + "\n".join(rows)
            sheets.append(text)
    except Exception as e:
        sheets = [f"[Error reading spreadsheet: {e}]"]
    return ParsedDocument(name=name, pages=sheets)


def _parse_docx(name: str, content: bytes) -> ParsedDocument:
    from docx import Document
    doc = Document(io.BytesIO(content))
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
    # Also grab table text
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if row_text:
                paragraphs.append(row_text)
    text = "\n".join(paragraphs) if paragraphs else "[empty docx]"
    return ParsedDocument(name=name, pages=[text])
